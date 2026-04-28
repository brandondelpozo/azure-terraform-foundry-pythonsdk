import azure.functions as func
import azure.durable_functions as df
import json
import logging
import os
import traceback
import importlib.util
from io import BytesIO
from datetime import datetime, timedelta, timezone
from typing import Dict, Any

from azure.storage.blob import (
    BlobServiceClient,
    generate_blob_sas,
    BlobSasPermissions,
    ContainerClient,
    ContentSettings,
)

# Import LangGraph pipeline instead of individual agents
from agents.langgraph_pipeline import run_langgraph_document_pipeline
from agents.parse_text_agent import AgentState
from agents.excel_langgraph_pipeline import run_excel_langgraph_pipeline

# Configure logging
logging.basicConfig(level=logging.INFO)

app = func.FunctionApp(http_auth_level=func.AuthLevel.ANONYMOUS)


def _load_pdf_converter():
    """Load pdf converter helper from utilities-functions/pdf_converter.py."""
    repo_root = os.path.dirname(os.path.dirname(__file__))
    converter_path = os.path.join(repo_root, "utilities-functions", "pdf_converter.py")

    spec = importlib.util.spec_from_file_location("pdf_converter", converter_path)
    if spec is None or spec.loader is None:
        raise ImportError(f"Unable to load pdf converter from {converter_path}")

    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module

# ---------------------------------------------------------------------------
# Helper: extract plain text from .txt or .docx bytes
# ---------------------------------------------------------------------------
def _extract_text(blob_bytes: bytes, filename: str) -> str:
    """Extract plain text from .txt or .docx bytes."""
    if filename.endswith(".txt"):
        return blob_bytes.decode("utf-8")
    elif filename.endswith(".docx"):
        import io
        from docx import Document
        logging.info("Parsing .docx content for %s", filename)
        doc = Document(io.BytesIO(blob_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        logging.info("Extracted %d non-empty paragraphs from %s", len(paragraphs), filename)
        return "\n".join(paragraphs)
    else:
        raise ValueError(f"Unsupported file type: {filename}")


# ---------------------------------------------------------------------------
# Helper: run LangGraph pipeline instead of manual chaining
# ---------------------------------------------------------------------------
def _run_agent_pipeline(text: str, filename: str = "document.docx") -> dict:
    """Run the LangGraph document processing pipeline with Completions API"""
    return run_langgraph_document_pipeline(text, filename)


def _build_docx_bytes_from_text(text: str) -> bytes:
    """Build a .docx document from plain text while preserving paragraph breaks."""
    from docx import Document

    document = Document()

    lines = text.split("\n")
    if not lines:
        lines = [text]

    for line in lines:
        document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# PATH A — Synchronous HTTP Trigger  POST /process-text
# ---------------------------------------------------------------------------
@app.route(route="process-text", methods=["POST"])
def http_start(req: func.HttpRequest) -> func.HttpResponse:
    """HTTP trigger that processes text with LangGraph pipeline."""
    try:
        req_body = req.get_json()
        if not req_body or "text" not in req_body:
            return func.HttpResponse(
                json.dumps({"error": "Missing 'text' in request body"}),
                status_code=400,
                mimetype="application/json",
            )

        text = req_body.get("text")
        filename = req_body.get("filename", "document.txt")
        
        logging.info(f"Processing text from /process-text: {text[:100]}...")

        # Use LangGraph pipeline
        final_state = _run_agent_pipeline(text, filename)
        
        logging.info(f"HTTP: Received final_state keys: {list(final_state.keys())}")
        logging.info(f"HTTP: final_state title: {final_state.get('title', 'MISSING')}")
        logging.info(f"HTTP: final_state summary: {final_state.get('summary', 'MISSING')}")
        logging.info(f"HTTP: final_state token_usage: {final_state.get('token_usage', 'MISSING')}")

        response = {
            "success": True,
            "original_text": final_state.get("text", ""),
            "enhanced_text": final_state.get("enhanced_text", ""),
            "title": final_state.get("title", ""),
            "summary": final_state.get("summary", ""),
            "synonyms_found": len(final_state.get("synonyms", {})),
            "synonyms_applied": final_state.get("synonyms", {}),
            "token_usage": final_state.get("token_usage", {}),
            "pdf_base64": final_state.get("pdf_content", ""),
            "parsed_data": final_state.get("parsed_data", {}),
            "workflow_info": {
                "agents_used": ["parse_text", "find_equivalents", "generate_summary", "consolidate_text"],
                "workflow_type": "langgraph_chat_completions_api_pipeline",
                "version": "3.2",  # Updated 2026-04-23 08:47
                "langgraph_enabled": True,
                "azure_openai_chat_completions_enabled": True,
            },
        }

        return func.HttpResponse(
            json.dumps(response, ensure_ascii=False),
            status_code=200,
            mimetype="application/json",
        )

    except Exception as e:
        logging.error(f"HTTP trigger error: {str(e)}")
        return func.HttpResponse(
            json.dumps({"error": str(e)}),
            status_code=500,
            mimetype="application/json",
        )


# ---------------------------------------------------------------------------
# PATH A.1 — Excel Processing HTTP Trigger POST /process-excel
# ---------------------------------------------------------------------------
@app.route(route="process-excel", methods=["POST"])
def process_excel_http(req: func.HttpRequest) -> func.HttpResponse:
    """HTTP trigger that processes Excel files with the LangGraph pipeline."""
    try:
        # Handle both multipart form data and JSON with base64 content
        content_type = req.headers.get('Content-Type', '')
        
        if 'multipart/form-data' in content_type:
            # Handle file upload
            files = req.files
            if 'file' not in files:
                return func.HttpResponse(
                    json.dumps({"error": "No file provided in multipart request"}),
                    status_code=400,
                    mimetype="application/json",
                )
            
            file = files['file']
            excel_content = file.read()
            filename = file.filename or "uploaded.xlsx"
            
        else:
            # Handle JSON with base64 content
            req_body = req.get_json()
            if not req_body or "excel_content" not in req_body:
                return func.HttpResponse(
                    json.dumps({"error": "Missing 'excel_content' in request body"}),
                    status_code=400,
                    mimetype="application/json",
                )

            import base64
            excel_content = base64.b64decode(req_body.get("excel_content"))
            filename = req_body.get("filename", "uploaded.xlsx")
        
        # Validate file extension
        if not filename.lower().endswith(('.xlsx', '.xls')):
            return func.HttpResponse(
                json.dumps({"error": "Only .xlsx and .xls files are supported"}),
                status_code=400,
                mimetype="application/json",
            )
        
        logging.info(f"Processing Excel file from /process-excel: {filename} ({len(excel_content)} bytes)")

        # Use Excel LangGraph pipeline
        result = run_excel_langgraph_pipeline(excel_content, filename)
        
        logging.info(f"Excel processing completed for: {filename}")
        logging.info(f"Generated title: {result.get('title', 'N/A')}")
        logging.info(f"Token usage: {result.get('token_usage', {})}")

        return func.HttpResponse(
            json.dumps(result, ensure_ascii=False),
            status_code=200,
            mimetype="application/json",
        )

    except Exception as e:
        logging.error(f"Excel processing HTTP trigger error: {str(e)}")
        return func.HttpResponse(
            json.dumps({"error": str(e)}),
            status_code=500,
            mimetype="application/json",
        )


# ---------------------------------------------------------------------------
# PATH B — Async Durable Orchestrator (Updated for LangGraph)
# ---------------------------------------------------------------------------
@app.orchestration_trigger(context_name="context")
def text_processing_orchestrator(context: df.DurableOrchestrationContext):
    """Orchestrator that calls the LangGraph pipeline"""
    try:
        input_data = context.get_input()
        text = input_data.get("text", "")
        filename = input_data.get("filename", "document.txt")

        logging.info(f"LangGraph orchestrator started for: {filename}")

        initial_state = {
            "text": text,
            "filename": filename
        }

        # Single activity call that runs the entire LangGraph
        final_state = yield context.call_activity("langgraph_pipeline_activity", initial_state)

        response = {
            "success": True,
            "original_text": final_state.get("text", ""),
            "enhanced_text": final_state.get("enhanced_text", ""),
            "title": final_state.get("title", ""),
            "summary": final_state.get("summary", ""),
            "synonyms_found": len(final_state.get("synonyms", {})),
            "synonyms_applied": final_state.get("synonyms", {}),
            "token_usage": final_state.get("token_usage", {}),
            "pdf_base64": final_state.get("pdf_content", ""),
            "parsed_data": final_state.get("parsed_data", {}),
            "workflow_info": {
                "agents_used": ["parse_text", "find_equivalents", "generate_summary", "consolidate_text"],
                "workflow_type": "langgraph_chat_completions_durable_functions",
                "version": "3.2",
                "instance_id": context.instance_id,
                "langgraph_enabled": True,
                "azure_openai_chat_completions_enabled": True,
            },
        }

        logging.info(f"LangGraph orchestration completed for instance: {context.instance_id}")
        return response

    except Exception as e:
        logging.error(f"LangGraph orchestrator error: {str(e)}")
        return {"error": str(e), "workflow": "langgraph_durable_functions"}


# Updated Activity Function for LangGraph
@app.activity_trigger(input_name="input")
def langgraph_pipeline_activity(input: dict) -> dict:
    """Activity that runs the entire LangGraph pipeline"""
    logging.info("Executing LangGraph pipeline activity")
    try:
        text = input.get("text", "")
        filename = input.get("filename", "document.txt")
        result = run_langgraph_document_pipeline(text, filename)
        logging.info("LangGraph pipeline activity completed successfully")
        return result
    except Exception as e:
        logging.error(f"LangGraph pipeline activity error: {str(e)}")
        raise


# Keep legacy activity functions for backward compatibility
@app.activity_trigger(input_name="input")
def parse_text_activity(input: dict) -> dict:
    logging.info("Executing legacy parse_text_activity - consider using langgraph_pipeline_activity")
    try:
        from agents.parse_text_agent import parse_text_agent
        result = parse_text_agent(input)
        logging.info("parse_text_activity completed successfully")
        return result
    except Exception as e:
        logging.error(f"parse_text_activity error: {str(e)}")
        raise


@app.activity_trigger(input_name="input")
def find_equivalents_activity(input: dict) -> dict:
    logging.info("Executing legacy find_equivalents_activity - consider using langgraph_pipeline_activity")
    try:
        from agents.find_equivalents_agent import find_equivalents_agent
        result = find_equivalents_agent(input)
        logging.info("find_equivalents_activity completed successfully")
        return result
    except Exception as e:
        logging.error(f"find_equivalents_activity error: {str(e)}")
        raise


@app.activity_trigger(input_name="input")
def consolidate_text_activity(input: dict) -> dict:
    logging.info("Executing legacy consolidate_text_activity - consider using langgraph_pipeline_activity")
    try:
        from agents.consolidate_text_agent import consolidate_text_agent
        result = consolidate_text_agent(input)
        logging.info("consolidate_text_activity completed successfully")
        return result
    except Exception as e:
        logging.error(f"consolidate_text_activity error: {str(e)}")
        raise


# ---------------------------------------------------------------------------
# PATH C — SAS URL generation  POST /generate-upload-url  (NEW)
# ---------------------------------------------------------------------------
@app.route(route="generate-upload-url", methods=["POST"])
def generate_upload_url(req: func.HttpRequest) -> func.HttpResponse:
    """
    Generates a write-only SAS URL so the client can upload files
    directly to Azure Blob Storage without needing Azure credentials.
    Supports .txt, .docx, .xlsx, and .xls files.
    """
    try:
        req_body = req.get_json()
        if not req_body or "filename" not in req_body:
            return func.HttpResponse(
                json.dumps({"error": "Missing 'filename' in request body"}),
                status_code=400,
                mimetype="application/json",
            )

        filename = req_body["filename"]
        ALLOWED = (".txt", ".docx", ".xlsx", ".xls")
        if not filename.endswith(ALLOWED):
            return func.HttpResponse(
                json.dumps({"error": "Only .txt, .docx, .xlsx, and .xls files are supported"}),
                status_code=400,
                mimetype="application/json",
            )

        account_name = os.environ["STORAGE_ACCOUNT_NAME"]
        account_key = os.environ["STORAGE_ACCOUNT_KEY"]
        container_name = "uploads"
        expiry_minutes = 15

        expiry = datetime.now(timezone.utc) + timedelta(minutes=expiry_minutes)

        sas_token = generate_blob_sas(
            account_name=account_name,
            container_name=container_name,
            blob_name=filename,
            account_key=account_key,
            permission=BlobSasPermissions(write=True, create=True),
            expiry=expiry,
        )

        sas_url = (
            f"https://{account_name}.blob.core.windows.net"
            f"/{container_name}/{filename}?{sas_token}"
        )

        logging.info(f"Generated SAS URL for blob: {filename}")

        return func.HttpResponse(
            json.dumps({
                "sas_url": sas_url,
                "blob_name": filename,
                "container": container_name,
                "expires_in_minutes": expiry_minutes,
            }),
            status_code=200,
            mimetype="application/json",
        )

    except Exception as e:
        logging.error(f"generate_upload_url error: {str(e)}")
        return func.HttpResponse(
            json.dumps({"error": str(e)}),
            status_code=500,
            mimetype="application/json",
        )


# ---------------------------------------------------------------------------
# PATH C — Blob Trigger (Updated for LangGraph + Azure OpenAI Completions)
# ---------------------------------------------------------------------------
@app.blob_trigger(
    arg_name="blob",
    path="uploads/{name}",
    connection="AzureWebJobsStorage",
    data_type="binary",
)
def process_blob(blob: func.InputStream):
    """
    LangGraph-powered document and Excel processing with Azure OpenAI Completions API
    Triggered by file upload to uploads/, saves an enhanced version to results/
    """
    blob_name = blob.name  # e.g. "uploads/mydoc.docx" or "uploads/myfile.xlsx"
    filename = blob_name.split("/")[-1]  # e.g. "mydoc.docx" or "myfile.xlsx"

    logging.info(f"LangGraph + Completions API processing started for: {blob_name}")
    
    account_name = os.environ["STORAGE_ACCOUNT_NAME"]
    account_key = os.environ["STORAGE_ACCOUNT_KEY"]
    base_name, _ = os.path.splitext(filename)
    result_blob_name = f"{base_name}.metadata.json"

    try:
        # Validate file type
        ALLOWED = (".txt", ".docx", ".xlsx", ".xls")
        if not filename.endswith(ALLOWED):
            logging.warning(f"Unsupported file type: {filename} - skipping")
            return

        # Read file content
        raw_bytes = blob.read()
        
        # Process based on file type
        if filename.lower().endswith(('.xlsx', '.xls')):
            # Excel file processing
            logging.info(f"Processing Excel file: {filename}")
            result = run_excel_langgraph_pipeline(raw_bytes, filename)
            
            # Prepare Excel-specific result metadata
            processing_result = {
                "success": result.get("success", False),
                "source_blob": blob_name,
                "metadata_blob": f"results/{result_blob_name}",
                "filename": filename,
                "title": result.get("title", ""),
                "summary": result.get("summary", ""),
#                "parsed_data": result.get("parsed_data", {})
                "token_usage": result.get("token_usage", {}),
                "workflow_info": result.get("workflow_info", {}),
                "processing_timestamp": datetime.now(timezone.utc).isoformat()
            }
            
        else:
            # Text/Word document processing (existing logic)
            text = _extract_text(raw_bytes, filename)
            logging.info(f"Extracted {len(text)} characters from: {filename}")

            if not text.strip():
                logging.warning(f"Empty document: {filename} - skipping")
                return

            # Run LangGraph pipeline with Completions API
            logging.info("Starting LangGraph pipeline with Azure OpenAI Completions...")
            final_state = _run_agent_pipeline(text, filename)
            
            # Prepare text document result metadata
            processing_result = {
                "success": True,
                "source_blob": blob_name,
                "result_blob": f"results/{base_name}_improved.docx",
                "metadata_blob": f"results/{result_blob_name}",
                "original_text": final_state.get("text", ""),
                "enhanced_text": final_state.get("enhanced_text", ""),
                "title": final_state.get("title", ""),
                "summary": final_state.get("summary", ""),
                "synonyms_found": len(final_state.get("synonyms", {})),
                "synonyms_applied": final_state.get("synonyms", {}),
                "token_usage": final_state.get("token_usage", {}),
                "parsed_data": final_state.get("parsed_data", {}),
                "workflow_info": {
                    "agents_used": ["parse_text", "find_equivalents", "generate_summary", "consolidate_text"],
                    "workflow_type": "langgraph_chat_completions_api_pipeline",
                    "version": "3.2",
                    "langgraph_enabled": True,
                    "azure_openai_chat_completions_enabled": True,
                    "processing_timestamp": datetime.now(timezone.utc).isoformat()
                },
            }

        # Save results to "results/" container
        blob_service = BlobServiceClient(
            account_url=f"https://{account_name}.blob.core.windows.net",
            credential=account_key,
        )
        results_container = blob_service.get_container_client("results")

        # For text documents, create and upload enhanced .docx
        if not filename.lower().endswith(('.xlsx', '.xls')):
            improved_docx_blob_name = f"{base_name}_improved.docx"
            enhanced_text = final_state.get("enhanced_text", "") or final_state.get("text", "")
            improved_docx_bytes = _build_docx_bytes_from_text(enhanced_text)
            
            # Upload enhanced .docx file
            results_container.upload_blob(
                name=improved_docx_blob_name,
                data=improved_docx_bytes,
                overwrite=True,
                content_settings=ContentSettings(
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ),
            )
            logging.info(f"Enhanced .docx saved to: results/{improved_docx_blob_name}")

        # Upload processing metadata
        results_container.upload_blob(
            name=result_blob_name,
            data=json.dumps(processing_result, ensure_ascii=False, indent=2),
            overwrite=True,
            content_settings=ContentSettings(content_type="application/json"),
        )
        logging.info(f"Processing metadata saved to: results/{result_blob_name}")

        logging.info(f"LangGraph + Completions processing completed for: {filename}")

    except Exception as e:
        # Error handling with detailed logging
        error_payload = {
            "success": False,
            "source_blob": blob_name,
            "error": str(e),
            "traceback": traceback.format_exc(),
            "workflow_info": {
                "workflow_type": "langgraph_completions_api_pipeline",
                "version": "3.2",
                "error_timestamp": datetime.now(timezone.utc).isoformat()
            },
        }

        logging.exception(f"LangGraph + Completions processing failed for: {blob_name}")

        # Save error details to results/
        try:
            blob_service = BlobServiceClient(
                account_url=f"https://{account_name}.blob.core.windows.net",
                credential=account_key,
            )
            results_container = blob_service.get_container_client("results")
            results_container.upload_blob(
                name=result_blob_name,
                data=json.dumps(error_payload, ensure_ascii=False, indent=2),
                overwrite=True,
            )
            logging.info(f"Error details saved to: results/{result_blob_name}")
        except Exception:
            logging.exception(f"Failed to save error details for: {blob_name}")
